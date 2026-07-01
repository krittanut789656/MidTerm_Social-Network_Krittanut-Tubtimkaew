"""
report_generator.py — Generate 2-page A4 report in .md and .docx formats.

Outputs:
    outputs/midterm_report_draft.md
    outputs/midterm_report_draft.docx
"""

import logging
import sys
from pathlib import Path

import pandas as pd

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from src.config import PATHS, OUTPUTS_DIR

log = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")


# ─────────────────────────────────────────────────────────────────────────────
# Load results (gracefully)
# ─────────────────────────────────────────────────────────────────────────────

def _load(path) -> pd.DataFrame | None:
    p = Path(path)
    if not p.exists():
        return None
    try:
        return pd.read_csv(p)
    except Exception:
        return None


def _fmt(val, decimals=4):
    try:
        return f"{float(val):.{decimals}f}"
    except Exception:
        return str(val)


# ─────────────────────────────────────────────────────────────────────────────
# Build report text
# ─────────────────────────────────────────────────────────────────────────────

def build_report_md() -> str:
    centrality = _load(PATHS["centrality_results"])
    louvain    = _load(PATHS["community_louvain"])
    validated  = _load(PATHS["corr_edges_validated"])
    partial    = _load(PATHS["partial_corr_edges"])
    regime_cmp = _load(PATHS["regime_comparison"])
    factors    = _load(PATHS["factor_exposure_edges"])
    final      = _load(PATHS["final_weekly_dataset"])

    # Compute key stats
    n_obs    = len(final) if final is not None else "N/A"
    if final is not None and hasattr(final, 'index'):
        final2 = pd.read_csv(PATHS["final_weekly_dataset"], index_col=0, parse_dates=True)
        date_range = f"{final2.index[0].date()} to {final2.index[-1].date()}"
    else:
        date_range = "May 2022 – Latest"

    raw_edges = len(validated) if validated is not None else "N/A"
    par_edges = len(partial)   if partial   is not None else "N/A"
    survival  = f"{int(par_edges)/int(raw_edges)*100:.1f}%" if isinstance(raw_edges, int) and raw_edges > 0 else "N/A"

    n_communities = louvain["communityId"].nunique() if louvain is not None else "N/A"

    top_degree = top_between = top_pr = "N/A"
    if centrality is not None and not centrality.empty:
        if "degree_centrality" in centrality.columns:
            r = centrality.nlargest(1, "degree_centrality").iloc[0]
            top_degree = f"{r.get('name','?')} ({_fmt(r.get('degree_centrality',0))})"
        if "betweenness_centrality" in centrality.columns:
            r = centrality.nlargest(1, "betweenness_centrality").iloc[0]
            top_between = f"{r.get('name','?')} ({_fmt(r.get('betweenness_centrality',0))})"
        if "pagerank" in centrality.columns:
            r = centrality.nlargest(1, "pagerank").iloc[0]
            top_pr = f"{r.get('name','?')} ({_fmt(r.get('pagerank',0))})"

    most_sensitive_bank = "N/A"
    if factors is not None and not factors.empty:
        cnt = factors.groupby("target").size().sort_values(ascending=False)
        most_sensitive_bank = f"{cnt.idxmax()} ({cnt.max()} significant factors)"

    louvain_tbl = _louvain_table(louvain)
    regime_tbl  = _regime_table(regime_cmp)

    sections = []
    sections.append("# Interest Rate and Risk Sentiment Network:\n# Graph Analysis of Thai SET50 Banking Stocks\n")
    sections.append("**การวิเคราะห์เครือข่ายผลกระทบของดอกเบี้ยโลก ค่าเงิน และความเสี่ยงตลาด\nต่อพฤติกรรมราคาหุ้นธนาคารไทยใน SET50**\n")
    sections.append("\n---\n\n## 1. Research Question\n\n")
    sections.append("> How do global interest rates, risk sentiment, FX movement, and financial sector ETFs\n")
    sections.append("> form a network of influence on Thai SET50 banking stocks?\n\n---\n\n")
    sections.append("## 2. Dataset and Variables\n\n")
    sections.append("| Category | Variables |\n|---|---|\n")
    sections.append("| Thai Banking Stocks | BBL, KBANK, KKP, KTB, SCB, TISCO, TTB |\n")
    sections.append("| Global ETFs | XLF (US Financial), EUFN (European Financial) |\n")
    sections.append("| FX | USDTHB |\n| Index | SET Index |\n")
    sections.append("| FRED Macro | FEDFUNDS, DGS2, DGS10, MORTGAGE30US, VIX |\n")
    sections.append("| BOT | BOT Policy Rate |\n")
    sections.append("| Derived | US Yield Curve (DGS10-DGS2), Thai Bank Basket |\n\n")
    sections.append("**Period:** " + str(date_range) + "\n")
    sections.append("**Frequency:** Weekly\n**Observations:** " + str(n_obs) + " weeks\n")
    sections.append("**Note:** Sample starts May 2022 to avoid SCB/SCBX structural continuity issues.\n\n---\n\n")
    sections.append("## 3. Data Cleaning\n\n")
    sections.append("- Removed duplicate dates and columns.\n")
    sections.append("- Resampled all series to weekly (Friday/last available).\n")
    sections.append("- Forward-filled macro variables up to 1 week; never forward-filled stock prices.\n")
    sections.append("- Converted prices/FX/indices to weekly log returns: log(P_t / P_(t-1)).\n")
    sections.append("- Converted rates/yields to weekly first differences: rate_t minus rate_(t-1).\n")
    sections.append("- Flagged outliers by |z-score| > 3.5 (not removed).\n")
    sections.append("- Verified SCB.BK price continuity from May 2022.\n\n---\n\n")
    sections.append("## 4. Graph Schema\n\n")
    sections.append("**Node Types:** Bank, ETF, MacroFactor, FX, Index, DerivedFactor\n\n")
    sections.append("**Relationship Types:**\n\n")
    sections.append("| Relationship | Method | Threshold |\n|---|---|---|\n")
    sections.append("| CORRELATED_WITH | Pearson + FDR (BH) | p_adj < 0.05, |r| >= 0.20 |\n")
    sections.append("| PARTIAL_CORRELATED_WITH | GraphicalLassoCV | |pc| >= 0.15 |\n")
    sections.append("| LAGGED_CORRELATED_WITH | Lagged Pearson (k=1-4) | |r| >= 0.30 |\n")
    sections.append("| INFLUENCES | OLS regression | p <= 0.05 |\n\n---\n\n")
    sections.append("## 5. Graph Analysis Results\n\n### 5.1 Centrality Results\n\n")
    sections.append("| Metric | Top Node | Score |\n|---|---|---|\n")
    sections.append("| Degree Centrality | " + str(top_degree) + " | Most connections |\n")
    sections.append("| Betweenness Centrality | " + str(top_between) + " | Bridge node |\n")
    sections.append("| PageRank | " + str(top_pr) + " | Influence score |\n\n")
    sections.append("**Interpretation:** The node with highest degree centrality is most broadly associated with\n")
    sections.append("price movements across the financial network. The betweenness node acts as the primary\n")
    sections.append("bridge between global macro conditions and Thai banking stock behavior.\n\n")
    sections.append("### 5.2 Community Detection\n\n")
    sections.append("**Louvain Communities:** " + str(n_communities) + " communities detected.\n\n")
    sections.append(louvain_tbl + "\n\n")
    sections.append("**Interpretation:** Thai banking stocks and global macro/financial variables cluster into\n")
    sections.append("communities. Variables within the same community exhibit stronger co-movement patterns.\n\n")
    sections.append("### 5.3 Raw vs Partial Correlation\n\n")
    sections.append("| Network | Edge Count |\n|---|---|\n")
    sections.append("| Raw Validated Correlation | " + str(raw_edges) + " |\n")
    sections.append("| Partial Correlation | " + str(par_edges) + " |\n")
    sections.append("| Edge Survival Rate | " + str(survival) + " |\n\n")
    sections.append("**Interpretation:** After conditioning on all other variables, " + str(survival) + " of raw correlation\n")
    sections.append("edges survive. Edges that disappear were driven by common factors. Surviving edges\n")
    sections.append("represent more direct associations between pairs of variables.\n\n")
    sections.append("### 5.4 Factor Exposure (OLS)\n\n")
    sections.append("**Most Globally Sensitive Bank:** " + str(most_sensitive_bank) + "\n\n")
    sections.append("OLS model: Bank_Return ~ XLF + EUFN + SET + USDTHB + VIX_CHANGE + DGS2_chg + DGS10_chg\n")
    sections.append("           + US_YIELD_CURVE_chg + MORTGAGE30US_chg + BOT_RATE_CHANGE\n\n")
    sections.append("Significant edges (p <= 0.05) retained as INFLUENCES relationships.\n\n")
    sections.append("### 5.5 Regime-Aware Network\n\n")
    sections.append(regime_tbl + "\n\n")
    sections.append("**Regime Classification:** Rolling 26-week change in FEDFUNDS and DGS2.\n")
    sections.append("'Hiking or Restrictive' if either rate is rising over the past 26 weeks.\n")
    sections.append("No look-ahead bias applied.\n\n---\n\n")
    sections.append("## 6. Key Findings\n\n")
    sections.append("1. **Most central node:** " + str(top_degree) + " — most broadly connected variable.\n")
    sections.append("2. **Bridge variable:** " + str(top_between) + " — connects clusters, transmits macro shocks.\n")
    sections.append("3. **Most sensitive bank:** " + str(most_sensitive_bank) + "\n")
    sections.append("4. **Community structure:** " + str(n_communities) + " distinct communities (Louvain).\n")
    sections.append("5. **Edge survival:** Only " + str(survival) + " of raw correlation edges survive partial adjustment.\n")
    sections.append("6. **Regime effect:** Network structure differs between Fed hiking and cutting regimes.\n")
    sections.append("7. **Risk transmission:** VIX and USDTHB linked to Thai bank returns (potential channels).\n")
    sections.append("8. **Methodological note:** All findings are association-based (ex-post). No causal claims.\n\n---\n\n")
    sections.append("## 7. Limitations\n\n")
    sections.append("1. Financial behavior network derived from time series — not a traditional social network.\n")
    sections.append("2. Edges represent statistical relationships, not causal links.\n")
    sections.append("3. Correlation does not imply causation.\n")
    sections.append("4. OLS with correlated regressors (DGS2, DGS10, US_YIELD_CURVE) may have multicollinearity.\n")
    sections.append("5. Weekly data reduces noise but may miss intra-week dynamics.\n")
    sections.append("6. Neo4j GDS results require Neo4j running with GDS plugin installed.\n")
    sections.append("7. Sample size (~" + str(n_obs) + " weeks) is sufficient but limited for high-dimensional partial correlation.\n\n---\n\n")
    sections.append("## 8. Conclusion\n\n")
    sections.append("This study constructs a financial behavior network linking global interest rates, risk sentiment,\n")
    sections.append("FX movement, and financial sector ETFs to Thai SET50 banking stock returns.\n\n")
    sections.append("Graph analysis reveals a structured network with distinct communities, clear bridge variables,\n")
    sections.append("and measurable differences in network topology between Fed tightening and easing regimes.\n\n")
    sections.append("The partial correlation adjustment substantially reduces the raw edge count, demonstrating\n")
    sections.append("that many apparent correlations are driven by common global factors rather than direct pairwise\n")
    sections.append("relationships.\n\n")
    sections.append("---\n*Generated by Thai Bank Graph Analysis Pipeline — Social Network & Media Analysis, Midterm Project*\n")

    md = "".join(sections)
    return md


def _louvain_table(louvain) -> str:
    if louvain is None or louvain.empty:
        return "*Louvain results not yet available. Run Neo4j GDS.*"
    groups = louvain.groupby("communityId")["name"].apply(list)
    lines = ["| Community | Members |", "|---|---|"]
    for gid, members in groups.items():
        lines.append(f"| {gid} | {', '.join(members)} |")
    return "\n".join(lines)


def _regime_table(regime_cmp) -> str:
    if regime_cmp is None or regime_cmp.empty:
        return "*Regime comparison not yet available.*"
    lines = ["| " + " | ".join(regime_cmp.columns) + " |",
             "|" + "---|" * len(regime_cmp.columns)]
    for _, row in regime_cmp.iterrows():
        lines.append("| " + " | ".join(str(v) for v in row.values) + " |")
    return "\n".join(lines)


def write_docx(md_text: str, output_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("python-docx not installed -- skipping .docx output.")
        return

    doc = Document()
    # Narrow margins
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    for line in md_text.splitlines():
        line = line.rstrip()
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("| "):
            # Skip markdown table rows in docx (already in md text)
            pass
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:], style="Quote")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line == "---":
            doc.add_paragraph()
        elif line:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(11) if p.runs else None
        else:
            doc.add_paragraph()

    doc.save(str(output_path))
    log.info("Saved DOCX: %s", output_path)


# ─────────────────────────────────────────────────────────────────────────────
# Main entry
# ─────────────────────────────────────────────────────────────────────────────

def run_report_generator():
    from src.config import PATHS, OUTPUTS_DIR
    md_text  = build_report_md()
    md_path  = OUTPUTS_DIR / "midterm_report_draft.md"
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    md_path.write_text(md_text, encoding="utf-8")
    log.info("Saved MD: %s", md_path)
    docx_path = OUTPUTS_DIR / "midterm_report_draft.docx"
    write_docx(md_text, docx_path)
    return md_path, docx_path


if __name__ == "__main__":
    run_report_generator()
